from __future__ import annotations

import csv
import hashlib
import io
import re
from dataclasses import dataclass
from pathlib import PurePosixPath
from typing import Any, Protocol
from zipfile import BadZipFile, ZipFile

from .models import Project, ResourceRecord, now_iso
from .object_store import ObjectStore
from .repository import InMemoryRepository

MAX_ATTACHMENT_BYTES = 20 * 1024 * 1024
MAX_EXTRACTED_CHARACTERS = 5_000_000
MAX_ARCHIVE_FILES = 5_000
MAX_ARCHIVE_UNCOMPRESSED_BYTES = 100 * 1024 * 1024
MAX_PDF_PAGES = 2_000
MAX_SHEETS = 200
MAX_ROWS_PER_SHEET = 100_000
MAX_COLUMNS_PER_SHEET = 2_000
MAX_CANDIDATES = 2_000
CHUNK_TARGET_CHARACTERS = 2_000
CHUNK_MAX_CHARACTERS = 4_000


class DocumentProcessingError(Exception):
    def __init__(self, code: str, safe_message: str) -> None:
        super().__init__(safe_message)
        self.code = code
        self.safe_message = safe_message


@dataclass(frozen=True, slots=True)
class ParsedUnit:
    text: str
    page: int | None = None
    sheet: str | None = None
    section: str | None = None
    row_start: int | None = None
    row_end: int | None = None

    def location(self) -> dict[str, Any]:
        return {
            key: value
            for key, value in {
                "page": self.page,
                "sheet": self.sheet,
                "section": self.section,
                "row_start": self.row_start,
                "row_end": self.row_end,
            }.items()
            if value is not None
        }


@dataclass(frozen=True, slots=True)
class ParsedDocument:
    format: str
    units: tuple[ParsedUnit, ...]


@dataclass(frozen=True, slots=True)
class DocumentChunk:
    id: str
    checksum_sha256: str
    ordinal: int
    text: str
    provenance: dict[str, Any]


class DocumentParser(Protocol):
    def parse(self, content: bytes) -> ParsedDocument: ...


def _clean_text(value: str) -> str:
    value = value.replace("\x00", " ").replace("\r\n", "\n").replace("\r", "\n")
    return "\n".join(" ".join(line.split()) for line in value.splitlines()).strip()


def _check_extracted_size(units: list[ParsedUnit]) -> None:
    if sum(len(unit.text) for unit in units) > MAX_EXTRACTED_CHARACTERS:
        raise DocumentProcessingError(
            "DOCUMENT_LIMIT_EXCEEDED", "Extracted document text exceeds the safety limit."
        )


def _validate_office_archive(content: bytes) -> None:
    try:
        with ZipFile(io.BytesIO(content)) as archive:
            entries = archive.infolist()
            if len(entries) > MAX_ARCHIVE_FILES:
                raise DocumentProcessingError(
                    "DOCUMENT_LIMIT_EXCEEDED", "Office document contains too many files."
                )
            total = 0
            for entry in entries:
                path = PurePosixPath(entry.filename.replace("\\", "/"))
                if path.is_absolute() or ".." in path.parts or entry.flag_bits & 0x1:
                    raise DocumentProcessingError(
                        "UNSAFE_DOCUMENT", "Office document contains an unsafe archive entry."
                    )
                total += entry.file_size
                if total > MAX_ARCHIVE_UNCOMPRESSED_BYTES:
                    raise DocumentProcessingError(
                        "DOCUMENT_LIMIT_EXCEEDED",
                        "Office document expands beyond the safety limit.",
                    )
                if entry.compress_size and entry.file_size / entry.compress_size > 1_000:
                    raise DocumentProcessingError(
                        "UNSAFE_DOCUMENT", "Office document has an unsafe compression ratio."
                    )
    except BadZipFile as exc:
        raise DocumentProcessingError("INVALID_DOCUMENT", "Office document is invalid.") from exc


class PdfParser:
    def parse(self, content: bytes) -> ParsedDocument:
        from pypdf import PdfReader

        try:
            reader = PdfReader(io.BytesIO(content), strict=True)
            if reader.is_encrypted:
                raise DocumentProcessingError(
                    "UNSUPPORTED_DOCUMENT", "Encrypted PDF documents are not supported."
                )
            if len(reader.pages) > MAX_PDF_PAGES:
                raise DocumentProcessingError(
                    "DOCUMENT_LIMIT_EXCEEDED", "PDF contains too many pages."
                )
            units = [
                ParsedUnit(text=text, page=index)
                for index, page in enumerate(reader.pages, start=1)
                if (text := _clean_text(page.extract_text() or ""))
            ]
        except DocumentProcessingError:
            raise
        except Exception as exc:
            raise DocumentProcessingError(
                "INVALID_DOCUMENT", "PDF document could not be parsed."
            ) from exc
        _check_extracted_size(units)
        return ParsedDocument("PDF", tuple(units))


class DocxParser:
    def parse(self, content: bytes) -> ParsedDocument:
        from docx import Document

        _validate_office_archive(content)
        try:
            document = Document(io.BytesIO(content))
            units: list[ParsedUnit] = []
            current_section = "Document"
            for paragraph in document.paragraphs:
                text = _clean_text(paragraph.text)
                if not text:
                    continue
                style = (paragraph.style.name if paragraph.style else "").lower()
                if style.startswith("heading"):
                    current_section = text
                units.append(ParsedUnit(text=text, section=current_section))
            for table_index, table in enumerate(document.tables, start=1):
                section = f"{current_section} / Table {table_index}"
                for row_index, row in enumerate(table.rows, start=1):
                    text = " | ".join(
                        cell_text for cell in row.cells if (cell_text := _clean_text(cell.text))
                    )
                    if text:
                        units.append(
                            ParsedUnit(
                                text=text,
                                section=section,
                                row_start=row_index,
                                row_end=row_index,
                            )
                        )
        except Exception as exc:
            raise DocumentProcessingError(
                "INVALID_DOCUMENT", "DOCX document could not be parsed."
            ) from exc
        _check_extracted_size(units)
        return ParsedDocument("DOCX", tuple(units))


class XlsxParser:
    def parse(self, content: bytes) -> ParsedDocument:
        from openpyxl import load_workbook

        _validate_office_archive(content)
        try:
            workbook = load_workbook(
                io.BytesIO(content), read_only=True, data_only=True, keep_links=False
            )
            if len(workbook.sheetnames) > MAX_SHEETS:
                raise DocumentProcessingError(
                    "DOCUMENT_LIMIT_EXCEEDED", "Workbook contains too many sheets."
                )
            units: list[ParsedUnit] = []
            for worksheet in workbook.worksheets:
                if (
                    worksheet.max_row > MAX_ROWS_PER_SHEET
                    or worksheet.max_column > MAX_COLUMNS_PER_SHEET
                ):
                    raise DocumentProcessingError(
                        "DOCUMENT_LIMIT_EXCEEDED", "Worksheet dimensions exceed the safety limit."
                    )
                for row_index, row in enumerate(worksheet.iter_rows(values_only=True), start=1):
                    values = [_clean_text(str(value)) for value in row if value is not None]
                    text = " | ".join(value for value in values if value)
                    if text:
                        units.append(
                            ParsedUnit(
                                text=text,
                                sheet=worksheet.title,
                                row_start=row_index,
                                row_end=row_index,
                            )
                        )
            workbook.close()
        except DocumentProcessingError:
            raise
        except Exception as exc:
            raise DocumentProcessingError(
                "INVALID_DOCUMENT", "XLSX document could not be parsed."
            ) from exc
        _check_extracted_size(units)
        return ParsedDocument("XLSX", tuple(units))


def _decode_text(content: bytes) -> str:
    try:
        return content.decode("utf-8-sig")
    except UnicodeDecodeError as exc:
        raise DocumentProcessingError(
            "INVALID_DOCUMENT", "Text documents must use UTF-8 encoding."
        ) from exc


class TxtParser:
    def parse(self, content: bytes) -> ParsedDocument:
        text = _decode_text(content)
        units = [
            ParsedUnit(text=cleaned, section=f"Paragraph {index}")
            for index, paragraph in enumerate(re.split(r"\n\s*\n", text), start=1)
            if (cleaned := _clean_text(paragraph))
        ]
        _check_extracted_size(units)
        return ParsedDocument("TXT", tuple(units))


class CsvParser:
    def parse(self, content: bytes) -> ParsedDocument:
        text = _decode_text(content)
        try:
            dialect = csv.Sniffer().sniff(text[:8_192], delimiters=",;\t|")
        except csv.Error:
            dialect = csv.excel
        units: list[ParsedUnit] = []
        try:
            for row_index, row in enumerate(csv.reader(io.StringIO(text), dialect), start=1):
                if row_index > MAX_ROWS_PER_SHEET or len(row) > MAX_COLUMNS_PER_SHEET:
                    raise DocumentProcessingError(
                        "DOCUMENT_LIMIT_EXCEEDED", "CSV dimensions exceed the safety limit."
                    )
                value = " | ".join(_clean_text(cell) for cell in row if _clean_text(cell))
                if value:
                    units.append(
                        ParsedUnit(
                            text=value,
                            sheet="CSV",
                            row_start=row_index,
                            row_end=row_index,
                        )
                    )
        except DocumentProcessingError:
            raise
        except csv.Error as exc:
            raise DocumentProcessingError("INVALID_DOCUMENT", "CSV document is invalid.") from exc
        _check_extracted_size(units)
        return ParsedDocument("CSV", tuple(units))


PARSERS: dict[str, DocumentParser] = {
    ".pdf": PdfParser(),
    ".docx": DocxParser(),
    ".xlsx": XlsxParser(),
    ".txt": TxtParser(),
    ".csv": CsvParser(),
}


def parse_document(filename: str, content: bytes) -> ParsedDocument:
    if len(content) > MAX_ATTACHMENT_BYTES:
        raise DocumentProcessingError(
            "DOCUMENT_LIMIT_EXCEEDED", "Attachment exceeds the 20 MiB safety limit."
        )
    suffix = "." + filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    parser = PARSERS.get(suffix)
    if parser is None:
        raise DocumentProcessingError("UNSUPPORTED_DOCUMENT", "Document type is not supported.")
    return parser.parse(content)


def _split_bounded(text: str) -> list[str]:
    if len(text) <= CHUNK_MAX_CHARACTERS:
        return [text]
    sentences = re.split(r"(?<=[.!?])\s+", text)
    pieces: list[str] = []
    current = ""
    for sentence in sentences:
        while len(sentence) > CHUNK_MAX_CHARACTERS:
            if current:
                pieces.append(current)
                current = ""
            pieces.append(sentence[:CHUNK_MAX_CHARACTERS])
            sentence = sentence[CHUNK_MAX_CHARACTERS:]
        candidate = f"{current} {sentence}".strip()
        if current and len(candidate) > CHUNK_TARGET_CHARACTERS:
            pieces.append(current)
            current = sentence
        else:
            current = candidate
    if current:
        pieces.append(current)
    return pieces


def chunk_document(document: ParsedDocument, attachment_checksum: str) -> tuple[DocumentChunk, ...]:
    chunks: list[DocumentChunk] = []
    ordinal = 0
    for unit_index, unit in enumerate(document.units, start=1):
        for part_index, text in enumerate(_split_bounded(unit.text), start=1):
            ordinal += 1
            provenance = {**unit.location(), "unit": unit_index, "part": part_index}
            checksum = hashlib.sha256(text.encode("utf-8")).hexdigest()
            stable = hashlib.sha256(
                f"{attachment_checksum}|{document.format}|{provenance}|{checksum}".encode()
            ).hexdigest()
            chunks.append(
                DocumentChunk(
                    id=f"dch_{stable[:32]}",
                    checksum_sha256=checksum,
                    ordinal=ordinal,
                    text=text,
                    provenance=provenance,
                )
            )
    return tuple(chunks)


def _candidate_sentences(text: str) -> list[str]:
    sentences = re.split(r"(?<=[.!?])\s+|\n+", text)
    return [sentence.strip() for sentence in sentences if 10 <= len(sentence.strip()) <= 2_000]


def _stable_resource_id(prefix: str, *values: str) -> str:
    digest = hashlib.sha256("|".join(values).encode("utf-8")).hexdigest()
    return f"{prefix}_{digest[:32]}"


def _find_resource(repository: InMemoryRepository, resource_id: str) -> ResourceRecord | None:
    return repository.resources.get(resource_id)


class DocumentIntelligenceService:
    """Parse and extract deterministically; no model input or document instruction is executed."""

    def __init__(self, repository: InMemoryRepository, object_store: ObjectStore) -> None:
        self.repository = repository
        self.object_store = object_store

    def process(self, project_id: str, job_id: str) -> None:
        try:
            with self.repository.transaction():
                project = self.repository.projects[project_id]
                job = self.repository.get_resource(project, job_id, "DOCUMENT_JOB")
                if job.status == "COMPLETED":
                    return
                job.status = "PROCESSING"
                job.data.update({"status": "PROCESSING", "progress": 10, "started_at": now_iso()})
                job.version += 1
                job.updated_at = now_iso()
                self.repository.persist()

            attachment = self.repository.get_resource(
                project, job.data["attachment_id"], "ATTACHMENT"
            )
            content = self.object_store.get(attachment.data["object_key"])
            if hashlib.sha256(content).hexdigest() != attachment.data["checksum_sha256"]:
                raise DocumentProcessingError(
                    "CHECKSUM_MISMATCH", "Stored attachment integrity verification failed."
                )
            parsed = parse_document(attachment.data["original_filename"], content)
            chunks = chunk_document(parsed, attachment.data["checksum_sha256"])

            with self.repository.transaction():
                project = self.repository.projects[project_id]
                job = self.repository.get_resource(project, job_id, "DOCUMENT_JOB")
                counts = self._persist_results(project, attachment, parsed, chunks, job)
                job.status = "COMPLETED"
                job.data.update(
                    {
                        "status": "COMPLETED",
                        "progress": 100,
                        "counts": counts,
                        "completed_at": now_iso(),
                        "error": None,
                    }
                )
                job.version += 1
                job.updated_at = now_iso()
                attachment.status = "ANALYZED"
                attachment.data["processing_status"] = "ANALYZED"
                attachment.data["semantically_analyzed"] = False
                attachment.data["deterministically_parsed"] = True
                attachment.version += 1
                attachment.updated_at = now_iso()
                self.repository.audit(
                    project.id,
                    "document_analysis_completed",
                    "SYSTEM",
                    {"job_id": job.id, "attachment_id": attachment.id, "counts": counts},
                )
                self.repository.persist()
        except DocumentProcessingError as exc:
            self._fail(project_id, job_id, exc.code, exc.safe_message)
        except (KeyError, FileNotFoundError):
            self._fail(
                project_id,
                job_id,
                "DOCUMENT_UNAVAILABLE",
                "The stored document is unavailable for processing.",
            )
        except Exception:  # noqa: BLE001 - boundary converts all failures to a safe job error
            self._fail(
                project_id,
                job_id,
                "DOCUMENT_PROCESSING_FAILED",
                "Document processing failed safely.",
            )

    def _fail(self, project_id: str, job_id: str, code: str, message: str) -> None:
        try:
            with self.repository.transaction():
                project = self.repository.projects[project_id]
                job = self.repository.get_resource(project, job_id, "DOCUMENT_JOB")
                job.status = "FAILED"
                job.data.update(
                    {
                        "status": "FAILED",
                        "progress": 100,
                        "failed_at": now_iso(),
                        "error": {"code": code, "message": message},
                    }
                )
                job.version += 1
                job.updated_at = now_iso()
                self.repository.audit(
                    project.id,
                    "document_analysis_failed",
                    "SYSTEM",
                    {"job_id": job.id, "error_code": code},
                )
                self.repository.persist()
        except Exception:  # noqa: BLE001 - failure reporting must not escape the worker boundary
            return

    def _persist_results(
        self,
        project: Project,
        attachment: ResourceRecord,
        parsed: ParsedDocument,
        chunks: tuple[DocumentChunk, ...],
        job: ResourceRecord,
    ) -> dict[str, int]:
        checksum = attachment.data["checksum_sha256"]
        source_id = _stable_resource_id("esrc", project.id, attachment.id, checksum)
        source = _find_resource(self.repository, source_id)
        if source is None:
            source = self.repository.create_resource(
                project,
                "EVIDENCE_SOURCE",
                {
                    "attachment_id": attachment.id,
                    "attachment_checksum_sha256": checksum,
                    "document_format": parsed.format,
                    "candidate_only": True,
                },
                status="CANDIDATE_UNCONFIRMED",
                provenance="DOCUMENT",
                immutable=True,
                resource_id=source_id,
            )

        parse_id = _stable_resource_id("dpr", project.id, attachment.id, checksum)
        if _find_resource(self.repository, parse_id) is None:
            self.repository.create_resource(
                project,
                "PARSED_DOCUMENT",
                {
                    "attachment_id": attachment.id,
                    "attachment_checksum_sha256": checksum,
                    "format": parsed.format,
                    "unit_count": len(parsed.units),
                    "chunk_count": len(chunks),
                    "parser_version": "LOCAL_DETERMINISTIC_V1",
                },
                status="COMPLETED",
                provenance="SYSTEM",
                immutable=True,
                resource_id=parse_id,
            )

        created_chunks = 0
        candidate_count = 0
        created_claims = 0
        created_evidence = 0
        for chunk in chunks:
            if _find_resource(self.repository, chunk.id) is None:
                self.repository.create_resource(
                    project,
                    "DOCUMENT_CHUNK",
                    {
                        "attachment_id": attachment.id,
                        "attachment_checksum_sha256": checksum,
                        "source_id": source.id,
                        "ordinal": chunk.ordinal,
                        "text": chunk.text,
                        "checksum_sha256": chunk.checksum_sha256,
                        "provenance": chunk.provenance,
                    },
                    status="PARSED",
                    provenance="DOCUMENT",
                    immutable=True,
                    resource_id=chunk.id,
                )
                created_chunks += 1
            for sentence_index, sentence in enumerate(_candidate_sentences(chunk.text), start=1):
                if candidate_count >= MAX_CANDIDATES:
                    break
                candidate_count += 1
                candidate_checksum = hashlib.sha256(sentence.encode("utf-8")).hexdigest()
                common = {
                    "attachment_id": attachment.id,
                    "attachment_checksum_sha256": checksum,
                    "source_id": source.id,
                    "chunk_id": chunk.id,
                    "chunk_checksum_sha256": chunk.checksum_sha256,
                    "candidate_checksum_sha256": candidate_checksum,
                    "candidate_index": sentence_index,
                    "provenance": chunk.provenance,
                    "confirmation": "UNCONFIRMED",
                    "extraction_method": "DETERMINISTIC_SENTENCE_SEGMENTATION",
                    "prompt_injection_treatment": "DOCUMENT_TEXT_ONLY",
                }
                claim_id = _stable_resource_id("clm", project.id, chunk.id, candidate_checksum)
                evidence_id = _stable_resource_id("evd", project.id, chunk.id, candidate_checksum)
                if _find_resource(self.repository, claim_id) is None:
                    self.repository.create_resource(
                        project,
                        "CLAIM",
                        {**common, "text": sentence},
                        status="CANDIDATE_UNCONFIRMED",
                        provenance="DOCUMENT",
                        immutable=False,
                        resource_id=claim_id,
                    )
                    project.memory["claims"].append(
                        {"id": claim_id, "status": "CANDIDATE_UNCONFIRMED"}
                    )
                    created_claims += 1
                if _find_resource(self.repository, evidence_id) is None:
                    self.repository.create_resource(
                        project,
                        "EVIDENCE",
                        {**common, "statement": sentence, "claim_id": claim_id},
                        status="CANDIDATE_UNCONFIRMED",
                        provenance="DOCUMENT",
                        immutable=False,
                        resource_id=evidence_id,
                    )
                    project.memory["evidence_ledger"].append(
                        {"id": evidence_id, "status": "CANDIDATE_UNCONFIRMED"}
                    )
                    created_evidence += 1
                if not any(
                    link["from_id"] == claim_id
                    and link["to_id"] == evidence_id
                    and link["relation"] == "CANDIDATE_SUPPORT"
                    for link in self.repository.creation_links
                ):
                    self.repository.add_link(
                        project, "CLAIM", claim_id, "EVIDENCE", evidence_id, "CANDIDATE_SUPPORT"
                    )

        self.repository.snapshot_memory(
            project,
            [
                {
                    "op": "ADD_CANDIDATE_DOCUMENT_EVIDENCE",
                    "attachment_id": attachment.id,
                    "job_id": job.id,
                }
            ],
        )
        return {
            "units": len(parsed.units),
            "chunks": len(chunks),
            "candidate_claims": candidate_count,
            "candidate_evidence": candidate_count,
            "new_chunks": created_chunks,
            "new_claims": created_claims,
            "new_evidence": created_evidence,
        }
